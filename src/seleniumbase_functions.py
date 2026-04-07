import base64
import pathlib
import re
from typing import Tuple, Optional, List, Union
import docx
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import StaleElementReferenceException, TimeoutException
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
import requests
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.action_chains import ActionChains
import pytesseract
from PIL import Image
import io
import os
from seleniumbase import Driver



class base(object):
    def __init__(self, is_headless_mode=False, timeout=5):
        # Đường dẫn đến thư mục Profile Chrome trên Mac của bạn
        # user_home = os.path.expanduser("~")
        # profile_path = f"{user_home}/Library/Application Support/Google/Chrome"
        
        # profile_name = "Profile 6"
        self.driver = Driver(
            uc=True,
            headless=is_headless_mode,
            # DÙNG PROFILE THẬT ĐỂ VƯỢT "ABOUT:BLANK"
            # user_data_dir=profile_path,
            # chromium_arg=f"--profile-directory={profile_name}",
            no_sandbox=True
        )
        self.driver.implicitly_wait(timeout)
        self.path = None
        self.timeout = timeout

    def create_folder(self, folder_name):
        cur_dir = os.path.abspath(os.getcwd())
        dir_name = folder_name
        self.path = os.path.join(cur_dir, dir_name)
        pathlib.Path(self.path).mkdir(parents=True, exist_ok=True)

    def maximize_browser(self):
        self.driver.maximize_window()

    def verify_element(self, element, timeOut=5):
        try:
            WebDriverWait(self.driver, timeOut).until(EC.visibility_of_element_located(element))
            return True
        except:
            return False

    def capture_screen(self, name):
        s = lambda x: self.driver.execute_script('return document.body.parentNode.scroll' + x)
        self.driver.set_window_size(s('Width'), s('Height'))  # May need manual adjustment
        self.driver.find_element(By.TAG_NAME, 'body').screenshot(name)

    def switch_frame(self, frame_reference):
        """Switch to a frame by WebElement, name, or index."""
        self.driver.switch_to.frame(frame_reference)

    def switch_back_to_default(self):
        self.driver.switch_to.default_content()

    @staticmethod
    def remove_file_if_exists(file):
        try:
            if os.path.exists(file):
                os.remove(file)
        except Exception:
            pass

    def define_img_name(self, title):
        file_name = os.path.join(self.path, title + ".png")
        return file_name

    def go_to_webpage(self, url):
        self.driver.get(url)

    def get_current_url(self):
        url = self.driver.current_url
        return url

    def click_element(self, element):
        if isinstance(element, WebElement):
            WebDriverWait(self.driver, self.timeout).until(EC.element_to_be_clickable(element)).click()
        elif isinstance(element, tuple):
            WebDriverWait(self.driver, self.timeout).until(EC.presence_of_element_located(element)).click()
        else:
            raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def reload_current_page(self):
        self.driver.refresh()

    def get_page_source(self) -> str:
        source = self.driver.page_source
        return source

    def wait_for_page_load(self, timeout: Optional[int] = None, poll_frequency: float = 0.5) -> bool:
        """Wait until document.readyState == 'complete'. Returns True if loaded, False on timeout."""
        wait_time = timeout if timeout is not None else self.timeout
        try:
            WebDriverWait(self.driver, wait_time, poll_frequency=poll_frequency).until(
                lambda d: d.execute_script('return document.readyState') == 'complete'
            )
            return True
        except TimeoutException:
            return False

    def wait_for_js_condition(self, js_condition: str, timeout: Optional[int] = None) -> bool:
        """Wait until the supplied JavaScript condition evaluates to a truthy value.

        Example: wait_for_js_condition('return window.someVar === true')
        """
        wait_time = timeout if timeout is not None else self.timeout
        try:
            WebDriverWait(self.driver, wait_time).until(
                lambda d: d.execute_script(js_condition)
            )
            return True
        except Exception:
            return False

    def wait_for_jquery(self, timeout: Optional[int] = None) -> bool:
        """Wait until jQuery active requests are finished (if jQuery is present).

        Falls back to True if jQuery is not present on the page.
        """
        wait_time = timeout if timeout is not None else self.timeout
        try:
            def _jq_inactive(d):
                try:
                    return d.execute_script('return (typeof jQuery !== "undefined") ? jQuery.active == 0 : true')
                except Exception:
                    return True

            WebDriverWait(self.driver, wait_time).until(_jq_inactive)
            return True
        except Exception:
            return False
    
    def wait_for_element_visible(self, element, timeout: Optional[int] = None) -> WebElement | None:
        """Wait until the element is visible and return it."""
        wait_time = timeout if timeout is not None else self.timeout
        if isinstance(element, tuple):
            try:
                return WebDriverWait(self.driver, wait_time).until(EC.visibility_of_element_located(element))
            except:
                return None
        elif isinstance(element, WebElement):
            try:
                return WebDriverWait(self.driver, wait_time).until(EC.visibility_of(element))
            except:
                return None
        raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def pass_data_to_file(self, source, file_name):
        try:
            f = open(file_name, "w")
            f.write(source)
        finally:
            f.close()
        
    def press_enter(self, element):
        if isinstance(element, WebElement):
            element.send_keys(Keys.ENTER)
        elif isinstance(element, tuple):
            self.find_element(element).send_keys(Keys.ENTER)
        else:
            raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def input_text(self, element, text):
        if isinstance(element, WebElement):
            element.clear()
            element.send_keys(text)
        elif isinstance(element, tuple):
            self.find_element(element).clear()
            self.find_element(element).send_keys(text)
        else:
            raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def open_new_tab(self, url):
        # Open a new blank tab and switch to the newly created window handle
        self.driver.execute_script("window.open('');")
        new_handle = self.driver.window_handles[-1]
        self.driver.switch_to.window(new_handle)
        self.driver.get(url)

    def switch_tab(self, tab_number):
        self.driver.switch_to.window(self.driver.window_handles[tab_number])

    def get_attribute_from_all_elements(self, elements, attribute_name):
        attrs = []
        list_of_elements = []
        if isinstance(elements, tuple):
            list_of_elements = self.find_elements(elements)
        elif isinstance(elements, list) and all(isinstance(e, WebElement) for e in elements):
            list_of_elements = elements
        else:
            raise TypeError("Invalid elements type. Must be a (By, str) tuple or a list of WebElements.")
        for item in list_of_elements:
            attrs.append(item.get_attribute(attribute_name))
        return attrs

    def get_attribute_from_element(self, element, attribute_name: str) -> str:
        if isinstance(element, WebElement):
            return element.get_attribute(attribute_name)
        elif isinstance(element, tuple):
            return self.find_element(element).get_attribute(attribute_name)
        raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    @staticmethod
    def replace_text(base_string, text_be_replaced, text_to_replace):
        new_string = str(base_string).replace(text_be_replaced, text_to_replace)
        return new_string

    @staticmethod
    def split_string(base_string, condition_split):
        data = str(base_string).split(condition_split)
        return data

    @staticmethod
    def crawl_data(url) -> BeautifulSoup | None:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        response = requests.get(url, headers=headers, timeout=20)
        if response.status_code == 200:
            soup = BeautifulSoup(response.content, "html.parser")
            return soup
        else:
            print(f"Page {url} returned status {response.status_code}. Skipping.")
            return None
    
    @staticmethod
    def crawl_text_from_soup(soup: BeautifulSoup, css_locator) -> str:
        for hidden_span in soup.find_all('span', style=lambda value: value and 'font-size:0' in value.replace(' ', '')):
            # Hàm decompose() sẽ xóa thẻ này và toàn bộ chữ rác bên trong nó khỏi soup
            hidden_span.decompose()
        element = soup.select_one(css_locator)
        if element:
            return element.get_text(separator='\n', strip=True)
        return ""


    def save_doc(self, title, body) -> None:
        document = docx.Document()
        try:
            document.add_heading(title.get_text())
            document.add_paragraph(body.get_text(separator='\n', strip=True))
            document.save(os.path.join(self.path, str(title.get_text()) + ".docx"))
        except:
            document.add_heading(title.text)
            if hasattr(body, 'get_text'):
                document.add_paragraph(body.get_text(separator='\n', strip=True))
            else:
                document.add_paragraph(body.text)
            document.save(os.path.join(self.path, str(title.text) + ".docx"))

    @staticmethod
    def get_element_by_tag(soup, tag, class_name):
        element = soup.find(tag, class_=class_name)
        return element

    @staticmethod
    def get_element_by_class(soup, class_name):
        element = soup.find(class_=class_name)
        return element

    def add_text_to_doc_file(self, title: str, text: str, file_name: Optional[str] = None):
        document = docx.Document()
        try:
            document.add_heading(title)
            
            # --- BẢN VÁ: TÁCH TỪNG DÒNG ĐỂ LƯU VÀO DOCX ---
            # Tránh lỗi python-docx dồn tất cả thành 1 cục và phá hủy dấu xuống dòng
            lines = text.split('\n')
            for line in lines:
                if line.strip():  # Chỉ thêm Paragraph nếu dòng đó thực sự có chữ
                    document.add_paragraph(line.strip())
            
            if file_name:
                document.save(os.path.join(self.path, file_name + ".docx"))
            else:
                document.save(os.path.join(self.path, title + ".docx"))
                
        except Exception as e:
            raise Exception(f"Exception while adding text to or saving doc file '{file_name or title}'.", e)
        
    @staticmethod
    def sleep(delay_time):
        time.sleep(delay_time)

    def scroll_web_page_to_the_end(self):
        last_height = self.driver.execute_script("return document.body.scrollHeight")

        while True:
            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            time.sleep(2)  # Adjust based on network/content load speed

            new_height = self.driver.execute_script("return document.body.scrollHeight")
            if new_height == last_height:
                break
            last_height = new_height

    def is_element_visible(self, element) -> bool:
        if isinstance(element, tuple):
            try:
                web_element = self.find_element(element)
            except:
                return False
        elif isinstance(element, WebElement):
            web_element = element
        else:
            raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

        return web_element.is_displayed()
    
    def action_click(self, element: Union[WebElement, Tuple[str, str]]) -> bool:
        """
        Clicks an element using simulated physical mouse movements.
        """
        try:
            target_element = element
            
            # If a locator tuple is passed, find the element first
            if isinstance(element, tuple):
                target_element = WebDriverWait(self.driver, self.timeout).until(
                    EC.presence_of_element_located(element)
                )
                
            if not isinstance(target_element, WebElement):
                raise TypeError("Input must be a (By, str) tuple or a WebElement.")

            # Scroll the element into the center of the viewport first
            self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", target_element)
            
            # Execute the human-like action chain
            actions = ActionChains(self.driver)
            actions.move_to_element(target_element).pause(0.3).click().perform()
            
            return True

        except Exception as e:
            print(f"Action click failed: {e}")
            return False

    def scroll_into_view(self, element):
        is_visible = self.is_element_visible(element)
        while not is_visible:
            self.driver.execute_script("arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", element)
            is_visible = self.is_element_visible(element)

    def find_element(self, element_, parent_element: Optional[WebElement] = None) -> Optional[WebElement]:
        target = parent_element if parent_element else self.driver
        if isinstance(element_, tuple):
            return target.find_element(*element_)
        elif isinstance(element_, WebElement):
            return element_ # Already a WebElement
        raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def find_elements(self, style_tuple: Tuple[By, str], parent_element: Optional[WebElement] = None) -> List[WebElement]:
        locator_strategy, locator_value = style_tuple
        if parent_element:
            elements = parent_element.find_elements(locator_strategy, locator_value)
        else:
            elements = self.driver.find_elements(locator_strategy, locator_value)
        return elements
        
    # def get_element_text(self, element) -> str:
    #     if isinstance(element, WebElement):
    #         return element.text.strip()
    #     elif isinstance(element, tuple):
    #         return self.find_element(element).text.strip()
    #     raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def get_element_text(self, element: Union[WebElement, Tuple[str, str]], extract_hidden: bool = False) -> str:
        try:
            target_element = element
            if isinstance(element, tuple):
                target_element = self.find_element(element)
                
            if not isinstance(target_element, WebElement):
                raise TypeError("Đầu vào bắt buộc phải là (By, str) tuple hoặc WebElement.")

            if extract_hidden:
                raw_text = target_element.get_attribute("textContent")
                return raw_text.strip() if raw_text else ""
            else:
                # ========================================================
                # BỘ QUÉT VĂN BẢN (TEXT SCANNER) MÔ PHỎNG MẮT NGƯỜI
                # ========================================================
                js_script = """
                var target = arguments[0];
                
                function extractText(node) {
                    // BƯỚC 1: XỬ LÝ TEXT NODE THUẦN TÚY (Chữ nằm giữa các thẻ)
                    if (node.nodeType === 3) { 
                        var parentStyle = window.getComputedStyle(node.parentNode);
                        // Bỏ qua chữ nếu thẻ cha đang tàng hình hoặc font-size = 0
                        if (parseFloat(parentStyle.fontSize) === 0 || 
                            parentStyle.display === 'none' || 
                            parentStyle.visibility === 'hidden' || 
                            parentStyle.opacity === '0' ||
                            parentStyle.color === 'rgba(0, 0, 0, 0)' ||
                            parentStyle.color === 'transparent') {
                            return "";
                        }
                        return node.nodeValue;
                    }
                    
                    // BƯỚC 2: XỬ LÝ ELEMENT NODE (Các thẻ div, p, span...)
                    if (node.nodeType === 1) { 
                        var style = window.getComputedStyle(node);
                        if (style.display === 'none' || style.visibility === 'hidden' || style.opacity === '0') {
                            return "";
                        }
                        
                        var text = "";
                        
                        // Xử lý thẻ xuống dòng <br> và thẻ khối (p, div) để tách đoạn
                        if (node.tagName === 'BR' || style.display === 'block' || node.tagName === 'P') {
                            text += "\\n";
                        }
                        
                        // LẤY CHỮ BỊ GIẤU TRONG CSS ::before
                        var before = window.getComputedStyle(node, '::before');
                        if (before && before.content && before.content !== 'none' && before.content !== 'normal') {
                            if (parseFloat(before.fontSize) > 0 && before.opacity !== '0') {
                                // Xóa dấu nháy kép (") bọc quanh chữ của CSS content
                                text += before.content.replace(/^["']|["']$/g, '');
                            }
                        }
                        
                        // Đi sâu vào đọc tiếp các thẻ con bên trong
                        for (var i = 0; i < node.childNodes.length; i++) {
                            text += extractText(node.childNodes[i]);
                        }
                        
                        // LẤY CHỮ BỊ GIẤU TRONG CSS ::after
                        var after = window.getComputedStyle(node, '::after');
                        if (after && after.content && after.content !== 'none' && after.content !== 'normal') {
                            if (parseFloat(after.fontSize) > 0 && after.opacity !== '0') {
                                text += after.content.replace(/^["']|["']$/g, '');
                            }
                        }
                        
                        return text;
                    }
                    return "";
                }
                
                var result = extractText(target);
                
                // BƯỚC 3: LÀM SẠCH VĂN BẢN
                // Xóa khoảng trắng thừa giữa các từ nhưng vẫn giữ nguyên các dấu xuống dòng
                result = result.replace(/\\r/g, '')
                               .replace(/[ \\t]+/g, ' ')
                               .replace(/\\n\\s+/g, '\\n')
                               .replace(/\\n{3,}/g, '\\n\\n');
                               
                return result.trim();
                """
                
                clean_text = self.driver.execute_script(js_script, target_element)
                return clean_text if clean_text else ""

        except StaleElementReferenceException:
            print("Cảnh báo: Element đã bị thay đổi (Stale).")
            return ""
        except Exception as e:
            print(f"Lỗi không xác định khi lấy text: {e}")
            return ""
    

    def close_browser(self):
        self.driver.close()

    def quit_driver(self):
        self.driver.quit()

    def get_content_fast(url, element, expected_attribute):
        # 1. Tải HTML tĩnh về 
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # 2. Tìm thẻ chứa chuỗi mã hóa bằng CSS Selector
        # Thay vì soup.find('div', id='chapterContentEncoded')
        # Dùng thẻ div kết hợp với dấu # để chỉ định ID
        encoded_div = soup.select_one(element) 
        
        # Mẹo nhỏ: Bạn có thể viết ngắn gọn hơn nữa là soup.select_one('#chapterContentEncoded') 
        
        if encoded_div and expected_attribute in encoded_div.attrs:
            encoded_string = encoded_div[expected_attribute]
            
            # 3. Dùng Python giải mã Base64 thành HTML gốc
            decoded_bytes = base64.b64decode(encoded_string)
            decoded_html = decoded_bytes.decode('utf-8')
            
            # 4. Lọc bỏ các thẻ HTML để lấy Text thuần
            clean_text = BeautifulSoup(decoded_html, 'html.parser').get_text(separator='\n')
            return clean_text
        else:
            return "No content found."
        
    def extract_images_to_file(self, elements: List[WebElement], file_name: str):
        """
        Cuộn đến từng ảnh trong danh sách truyền vào, quét OCR và lưu nối tiếp vào file.
        
        Args:
            elements (List[WebElement]): Danh sách các thẻ chứa ảnh cần quét.
            file_path (str): Đường dẫn/Tên file text để lưu kết quả.
        """
        if not elements:
            print("Danh sách thẻ (elements) truyền vào đang trống! Không có gì để quét.")
            return

        print(f"Bắt đầu quét {len(elements)} bức ảnh và lưu vào: {file_name}")
        file_path = os.path.join(self.path, file_name + ".docx")

        # Mở file ở chế độ "a" (Append)
        with open(file_path, "a", encoding="utf-8") as file:
            
            for index, img in enumerate(elements, start=1):
                try:
                    # Cuộn màn hình đến vị trí bức ảnh
                    self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", img)
                    time.sleep(1.5) # Đợi Lazy-load
                    
                    # Chụp ảnh và mở bằng Pillow
                    image_bytes = img.screenshot_as_png
                    image = Image.open(io.BytesIO(image_bytes))
                    
                    # Gọi Tesseract
                    print(f"Đang dịch ảnh {index}/{len(elements)}...")
                    text = pytesseract.image_to_string(image, lang='vie')
                    
                    # Ghi vào file
                    if text.strip():
                        file.write(text.strip() + "\n\n")
                        file.flush() 
                        
                except Exception as e:
                    print(f"Lỗi ở ảnh thứ {index}: {e}")
                    file.write(f"\n[Lỗi không đọc được ảnh {index}]\n\n")

        print("\nHOÀN THÀNH! Toàn bộ nội dung đã được lưu.")
