import pathlib
import re
from typing import Literal, Tuple, Optional, List, Union
import docx
import time
from docx.shared import Inches
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.remote.webelement import WebElement
from selenium.webdriver.support.ui import WebDriverWait
from selenium.common.exceptions import ElementClickInterceptedException, NoSuchElementException, StaleElementReferenceException, TimeoutException
from selenium.webdriver.support import expected_conditions as EC
from bs4 import BeautifulSoup
import requests
from selenium.webdriver.common.keys import Keys
import pytesseract
from curl_cffi import requests
from PIL import Image
import io
import os
from seleniumbase import Driver


class base(object):

    def __init__(self, is_headless_mode=True, timeout=5, use_seleniumbase=False, page_load_strategy="eager"):
        """Initializes the browser with configuration options."""
        self.path = None
        self.timeout = timeout

        if use_seleniumbase:
            # ---------------------------------------------------------
            # SELENIUMBASE (Bypass Cloudflare, Anti-bot)
            # ---------------------------------------------------------
            print(f"Starting browser with SeleniumBase (UC Mode, strategy={page_load_strategy})...")
            self.driver = Driver(
                uc=True,
                headless=is_headless_mode,
                no_sandbox=True,
                page_load_strategy=page_load_strategy,
                window_size="1920,1080"
               
            )
        else:
            # ---------------------------------------------------------
            # SELENIUM
            # ---------------------------------------------------------
            print(f"Starting browser with Selenium (strategy={page_load_strategy})...")
            options = webdriver.ChromeOptions()
            options.add_argument('--deny-permission-prompts')
            if is_headless_mode:
                options.add_argument('--headless=new')
            options.page_load_strategy = page_load_strategy
            options.add_argument('--disable-dev-shm-usage') 
            options.add_argument('--no-sandbox')
            options.add_argument('--window-size=1920,1080')
                
            self.driver = webdriver.Chrome(options=options)
        
        # Set Implicit Wait for element discovery
        self.driver.implicitly_wait(timeout)
        
        # Set page load timeout to 20 seconds
        self.driver.set_page_load_timeout(20)


    def set_path(self, folder_name):
        """Specifies the working directory without creating it (specifically for multi-threading)."""
        cur_dir = os.path.abspath(os.getcwd())
        self.path = os.path.join(cur_dir, folder_name)


    def create_folder(self, folder_name):
        """Creates a new folder and sets it as the current working path."""
        cur_dir = os.path.abspath(os.getcwd())
        dir_name = folder_name
        self.path = os.path.join(cur_dir, dir_name)
        pathlib.Path(self.path).mkdir(parents=True, exist_ok=True)

    def maximize_browser(self):
        """Maximizes the browser window."""
        self.driver.maximize_window()

    def verify_element(self, element, timeOut=5):
        """Checks if an element is visible on the page within the specified timeout."""
        try:
            WebDriverWait(self.driver, timeOut).until(EC.visibility_of_element_located(element))
            return True
        except:
            return False

    def capture_full_screen(self, file_name: str):
        """
        Safely captures a full-page screenshot from top to bottom.
        """
        full_file_path = os.path.join(self.path, f"{file_name}.png")
        print(f"Capturing full screen to: {full_file_path}")
        
        # STEP 1: HANDLE LAZY-LOAD (Mandatory)
        # Scroll slowly to the bottom to force the page to load all hidden images/content
        self.scroll_web_page_to_the_end(pause_time=1, max_scrolls=20)
        
        # Scroll back to top to prepare for capture
        self.driver.execute_script("window.scrollTo(0, 0);")
        time.sleep(1) # Wait for sticky headers/menus to collapse

        try:
            # STEP 2: CAPTURE FULL PAGE BASED ON DRIVER TYPE
            
            # Using SeleniumBase (Most robust, no OS limitations)
            if hasattr(self.driver, "save_page_screenshot"):
                self.driver.save_page_screenshot(full_file_path)
                print("Completed using SeleniumBase!")
                
            # Using standard Selenium (via Chrome DevTools Protocol - CDP)
            else:
                # Get actual layout metrics using CDP (bypasses viewport limitations)
                metrics = self.driver.execute_cdp_cmd('Page.getLayoutMetrics', {})
                content_width = metrics['contentSize']['width']
                content_height = metrics['contentSize']['height']
                
                # Send core-level screenshot command to Chrome
                screenshot_dict = self.driver.execute_cdp_cmd(
                    'Page.captureScreenshot', {
                        'format': 'png',
                        'captureBeyondViewport': True, # Capture beyond the viewport
                        'clip': {
                            'width': content_width,
                            'height': content_height,
                            'x': 0,
                            'y': 0,
                            'scale': 1
                        }
                    }
                )
                
                # Save image file from Base64 data
                import base64
                with open(full_file_path, "wb") as f:
                    f.write(base64.b64decode(screenshot_dict['data']))
                print("Completed with Chrome CDP!")
                
        except Exception as e:
            print(f"An error occurred while capturing the full screen: {e}")


    def switch_frame(self, frame_reference):
        """Switch to a frame by WebElement, name, or index."""
        self.driver.switch_to.frame(frame_reference)

    def switch_back_to_default(self):
        """Switches back to the main content from an iframe."""
        self.driver.switch_to.default_content()

    @staticmethod
    def remove_file_if_exists(file):
        """Removes a file from the disk if it exists."""
        try:
            if os.path.exists(file):
                os.remove(file)
        except Exception:
            pass

    def define_img_name(self, title):
        """Defines the full path for an image file based on the title."""
        file_name = os.path.join(self.path, title + ".png")
        return file_name

    def go_to_webpage(self, url: str, bypass_cloudflare: bool = False, reconnect_time: int = 4):
        """
        Navigates to a webpage.
        If bypass_cloudflare = True, uses UC Mode to bypass protection.
        If bypass_cloudflare = False, uses standard get() (faster).
        """
        
        try:
            # Use UC reconnection if enabled and using SeleniumBase
            if bypass_cloudflare and hasattr(self.driver, "uc_open_with_reconnect"):
                self.driver.uc_open_with_reconnect(url, reconnect_time=reconnect_time)
                
            else:
                # High-speed access using existing cookies
                self.driver.get(url)
                
        except Exception as e:
            print(f"Error while loading {url}: {e}")

    def get_current_url(self):
        """Returns the current URL of the webpage."""
        url = self.driver.current_url
        return url

    def click_element(self, element: Union[WebElement, Tuple[str, str]], force_js: bool = False) -> bool:
        """Clicks an element, optionally using JavaScript if a standard click fails."""
        try:
            target_element = element
            
            if isinstance(element, tuple):
                target_element = WebDriverWait(self.driver, self.timeout).until(
                    EC.presence_of_element_located(element)
                )
                
            if not isinstance(target_element, WebElement):
                raise TypeError("Element should be (By, str) tuple or WebElement.")

            self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", target_element)

            if force_js:
                self.driver.execute_script("arguments[0].click();", target_element)
            else:
                WebDriverWait(self.driver, self.timeout).until(
                    EC.element_to_be_clickable(target_element)
                )
                target_element.click()
                
            return True

        # (FALLBACK)
        except ElementClickInterceptedException:
            print("Element is not interactable. Using JS Click to rescue...")
            self.driver.execute_script("arguments[0].click();", target_element)
            return True
            
        except TimeoutException:
            print(f"Error: Unable to find or click element within {self.timeout} seconds")
            return False
            
        except StaleElementReferenceException:
            print("Warning: Element has been stale.")
            return False
            
        except Exception as e:
            print(f"Unknown error: {e}")
            return False

    def reload_current_page(self):
        """Reloads the current page."""
        self.driver.refresh()

    def get_page_source(self) -> str:
        """Returns the HTML source code of the current page."""
        source = self.driver.page_source
        return source

    def wait_for_page_load(self, timeout: Optional[int] = None, poll_frequency: float = 0.5) -> bool:
        """Wait until document.readyState == 'complete'. Returns True if loaded, False on timeout."""
        wait_time = timeout if timeout is not None else self.timeout
        try:
            WebDriverWait(self.driver, wait_time, poll_frequency=poll_frequency).until(
                lambda d: d.execute_script('return document.readyState') == 'complete'
            )
            return True
        except TimeoutException:
            return False

    def wait_for_js_condition(self, js_condition: str, timeout: Optional[int] = None) -> bool:
        """Wait until the supplied JavaScript condition evaluates to a truthy value.

        Example: wait_for_js_condition('return window.someVar === true')
        """
        wait_time = timeout if timeout is not None else self.timeout
        try:
            WebDriverWait(self.driver, wait_time).until(
                lambda d: d.execute_script(js_condition)
            )
            return True
        except Exception:
            return False

    def wait_for_jquery(self, timeout: Optional[int] = None) -> bool:
        """Wait until jQuery active requests are finished (if jQuery is present).

        Falls back to True if jQuery is not present on the page.
        """
        wait_time = timeout if timeout is not None else self.timeout
        try:
            def _jq_inactive(d):
                try:
                    return d.execute_script('return (typeof jQuery !== "undefined") ? jQuery.active == 0 : true')
                except Exception:
                    return True

            WebDriverWait(self.driver, wait_time).until(_jq_inactive)
            return True
        except Exception:
            return False
    
    def wait_for_element_visible(self, element, timeout: Optional[int] = None) -> bool:
        """Wait until the element is visible and return it."""
        wait_time = timeout if timeout is not None else self.timeout
        if isinstance(element, tuple):
            try:
                WebDriverWait(self.driver, wait_time).until(EC.visibility_of_element_located(element))
                return True
            except:
                return False
        elif isinstance(element, WebElement):
            try:
                WebDriverWait(self.driver, wait_time).until(EC.visibility_of(element))
                return True
            except:
                return False
        raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def pass_data_to_file(self, source, file_name):
        """Writes data (string) to a file."""
        try:
            f = open(file_name, "w")
            f.write(source)
        finally:
            f.close()
        
    def press_enter(self, element):
        """Sends an ENTER key to an element."""
        if isinstance(element, WebElement):
            element.send_keys(Keys.ENTER)
        elif isinstance(element, tuple):
            self.find_element(element).send_keys(Keys.ENTER)
        else:
            raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def input_text(self, element, text):
        """Clears existing content and types new text into an element."""
        if isinstance(element, WebElement):
            element.clear()
            element.send_keys(text)
        elif isinstance(element, tuple):
            self.find_element(element).clear()
            self.find_element(element).send_keys(text)
        else:
            raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def open_new_tab(self, url):
        """Opens a new tab and navigates to the specified URL."""
        # Open a new blank tab and switch to the newly created window handle
        self.driver.execute_script("window.open('');")
        new_handle = self.driver.window_handles[-1]
        self.driver.switch_to.window(new_handle)
        self.driver.get(url)

    def switch_tab(self, tab_number):
        """Switches to a tab by its index."""
        self.driver.switch_to.window(self.driver.window_handles[tab_number])

    def get_attribute_from_all_elements(self, elements, attribute_name):
        """Retrieves a specific attribute from all elements in a list."""
        attrs = []
        list_of_elements = []
        if isinstance(elements, tuple):
            list_of_elements = self.find_elements(elements)
        elif isinstance(elements, list) and all(isinstance(e, WebElement) for e in elements):
            list_of_elements = elements
        else:
            raise TypeError("Invalid elements type. Must be a (By, str) tuple or a list of WebElements.")
        for item in list_of_elements:
            attrs.append(item.get_attribute(attribute_name))
        return attrs

    def get_attribute_from_element(self, element, attribute_name: str) -> str:
        """Retrieves a specific attribute from an element."""
        if isinstance(element, WebElement):
            return element.get_attribute(attribute_name)
        elif isinstance(element, tuple):
            return self.find_element(element).get_attribute(attribute_name)
        raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    @staticmethod
    def replace_text(base_string, text_be_replaced, text_to_replace):
        """Replaces text within a string."""
        new_string = str(base_string).replace(text_be_replaced, text_to_replace)
        return new_string

    @staticmethod
    def split_string(base_string, condition_split):
        """Splits a string based on a condition."""
        data = str(base_string).split(condition_split)
        return data

    @staticmethod
    def crawl_data(
        url: str, 
        return_type: Literal["soup", "text", "content"] = "soup",
        impersonate: Optional[str] = "chrome120"
    ) -> Union[BeautifulSoup, str, bytes, None]:
        """
        Crawls a URL and returns the data in the specified format.
        
        :param url: The target URL.
        :param return_type: "soup" (BeautifulSoup obj), "text" (HTML string), or "content" (raw bytes).
        :param impersonate: Browser to impersonate (e.g., "chrome120", "safari15_3"). Set to None for a standard request.
        """
        try:
            # Set up request arguments dynamically
            request_args = {"timeout": 20}
            
            if impersonate:
                request_args["impersonate"] = impersonate
            else:
                # Fallback headers if we aren't impersonating a browser
                request_args["headers"] = {
                    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
                }

            # Execute the request
            response = requests.get(url, **request_args)

            if response.status_code == 200:
                # Return the data based on the requested type
                if return_type == "soup":
                    return BeautifulSoup(response.content, "html.parser")
                elif return_type == "text":
                    return response.text
                elif return_type == "content":
                    return response.content
                else:
                    print(f"Invalid return_type '{return_type}'. Returning None.")
                    return None
            else:
                print(f"Page {url} returned status {response.status_code}. Skipping.")
                return None
                
        except Exception as e:
            print(f"Error crawling {url}: {e}")
            return None


    @staticmethod
    def crawl_text_from_soup(data: Union[BeautifulSoup, str, None], css_locator: str) -> str:
        """Extracts text and preserves image positions using placeholders from Soup or HTML string."""
        
        # 1. Safely handle if crawl_data failed and returned None
        if not data:
            print("No data provided to extract_text. Skipping.")
            return ""
            
        # 2. If data is a raw HTML string, convert it to Soup automatically
        if isinstance(data, str):
            soup = BeautifulSoup(data, "html.parser")
        else:
            soup = data # It's already a BeautifulSoup object

        # 3. Find the element
        element = soup.select_one(css_locator)
        if not element:
            print(f"Could not find element with locator: {css_locator}")
            return ""

        # --- The rest of the original logic ---
        
        # Remove hidden noise (Anti-scraping)
        for hidden_span in element.find_all('span', style=lambda value: value and 'font-size:0' in value.replace(' ', '')):
            hidden_span.decompose()

        # FIND AND MARK IMAGES
        for img in element.find_all('img'):
            src = img.get('src') or img.get('data-src') or img.get('data-original')
            if src:
                if src.startswith("//"):
                    src = "https:" + src
                placeholder = f"\n[IMAGE_MARKER_START]{src}[IMAGE_MARKER_END]\n"
                img.replace_with(placeholder)

        # Get full text
        return element.get_text(separator='\n', strip=True)


    def save_doc(self, title, body) -> None:
        """Saves a title and body to a Word file (.docx)."""
        document = docx.Document()
        try:
            document.add_heading(title.get_text())
            document.add_paragraph(body.get_text(separator='\n', strip=True))
            document.save(os.path.join(self.path, str(title.get_text()) + ".docx"))
        except:
            document.add_heading(title.text)
            if hasattr(body, 'get_text'):
                document.add_paragraph(body.get_text(separator='\n', strip=True))
            else:
                document.add_paragraph(body.text)
            document.save(os.path.join(self.path, str(title.text) + ".docx"))

    @staticmethod
    def get_element_by_tag(soup, tag, class_name):
        """Finds an element based on tag and class name using BeautifulSoup."""
        element = soup.find(tag, class_=class_name)
        return element

    @staticmethod
    def get_element_by_class(soup, class_name):
        """Finds an element based on class name using BeautifulSoup."""
        element = soup.find(class_=class_name)
        return element

    def add_text_to_doc_file(self, title: str, text: str, file_name: Optional[str] = None):
        """Adds text (with image support via markers) to a new Word file."""
        document = docx.Document()
        try:
            # Add title
            document.add_heading(title)
            
            # Split content into lines
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    document.add_paragraph("") # Maintain spacing
                    continue
                
                # CHECK: IS THIS LINE AN IMAGE OR TEXT?
                if "[IMAGE_MARKER_START]" in line:
                    # Extract image link from marker using Regex
                    # (Handles cases where HTML might mix text and images on one line)
                    parts = re.split(r'\[IMAGE_MARKER_START\](.*?)\[IMAGE_MARKER_END\]', line)
                    
                    for part in parts:
                        part = part.strip()
                        if not part: continue
                        
                        # If image link -> Download and insert image
                        if part.startswith("http"):
                            try:
                                res = requests.get(part, timeout=10)
                                if res.status_code == 200:
                                    image_stream = io.BytesIO(res.content)
                                    document.add_picture(image_stream, width=Inches(5.0))
                                else:
                                    document.add_paragraph(f"[Could not download illustration: Error {res.status_code}]")
                            except Exception:
                                document.add_paragraph(f"[Connection lost while downloading illustration]")
                        
                        # If accidental accompanying text -> Write as normal text
                        else:
                            document.add_paragraph(part)
                            
                # If normal text line
                else:
                    document.add_paragraph(line)
            
            # Save file safely
            safe_name = file_name if file_name else re.sub(r'[\\/*?:"<>|]', "", title).strip()
            final_path = os.path.join(self.path, safe_name + ".docx")
            document.save(final_path)
            
        except Exception as e:
            raise Exception(f"Error while writing content to Word file '{file_name or title}': {e}")
        
    @staticmethod
    def sleep(delay_time):
        """Pauses execution for the specified duration."""
        time.sleep(delay_time)

    def scroll_web_page_to_the_end(self, pause_time=2, max_scrolls=50) -> bool:
        """
        Scrolls to the bottom of the page.
        Returns True if it hit the absolute bottom, False if it hit the max_scrolls limit.
        """
        last_height = self.driver.execute_script("return document.body.scrollHeight")
        scroll_count = 0

        while scroll_count < max_scrolls:
            # Scroll down to bottom
            self.driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
            
            # Wait to load page
            time.sleep(pause_time)

            # Calculate new scroll height and compare with last scroll height
            new_height = self.driver.execute_script("return document.body.scrollHeight")
            
            if new_height == last_height:
                # DOUBLE CHECK: Give it one more second to be absolutely sure it's not just a slow network
                time.sleep(1)
                new_height_check = self.driver.execute_script("return document.body.scrollHeight")
                if new_height_check == last_height:
                    print(f"Reached absolute bottom after {scroll_count} scrolls.")
                    return True # Signal: Success
                    
            last_height = new_height
            scroll_count += 1

        print(f"Warning: Stopped scrolling after hitting the maximum limit of {max_scrolls} scrolls.")
        return False # Signal: Reached limit before finding the bottom

    def is_element_visible(self, element) -> bool:
        """Checks if an element is currently displayed."""
        if isinstance(element, tuple):
            try:
                web_element = self.find_element(element)
                return web_element.is_displayed()
            except:
                return False
        elif isinstance(element, WebElement):
            try:
                web_element = element
                return web_element.is_displayed()
            except:
                return False
        else:
            raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")


    def scroll_into_view(self, element) -> bool:
        """
        Scrolls the page up or down to center the element.
        Returns True if the element becomes visible, False otherwise.
        """
        try:
            # 1. Execute the scroll command EXACTLY ONCE
            self.driver.execute_script(
                "arguments[0].scrollIntoView({behavior: 'smooth', block: 'center'});", 
                element
            )
            
            # 2. Pause briefly to allow the smooth scrolling animation to finish
            time.sleep(0.8) # Adjust this if the page is very long/slow
            
            # 3. Check visibility once after arriving
            if self.is_element_visible(element):
                return True
            else:
                print("Warning: Scrolled to element, but it is still hidden (e.g., behind a banner or display: none).")
                return False
                
        except Exception as e:
            print(f"Failed to scroll to element: {e}")
            return False

    def find_element(self, element_, parent_element: Optional[WebElement] = None) -> Optional[WebElement]:
        """Finds a single element based on a locator."""
        target = parent_element if parent_element else self.driver
        
        if isinstance(element_, tuple):
            try:
                # Try to find it
                return target.find_element(*element_)
            except NoSuchElementException:
                # If it fails, silently return None
                return None 
                
        elif isinstance(element_, WebElement):
            return element_
            
        raise TypeError("Invalid element type. Must be a (By, str) tuple or a WebElement.")

    def find_elements(self, style_tuple: Tuple[By, str], parent_element: Optional[WebElement] = None) -> List[WebElement]:
        """Finds all elements matching a locator."""
        locator_strategy, locator_value = style_tuple
        if parent_element:
            elements = parent_element.find_elements(locator_strategy, locator_value)
        else:
            elements = self.driver.find_elements(locator_strategy, locator_value)
        return elements

    def get_element_text(self, element: Union[WebElement, Tuple[str, str]], extract_hidden: bool = False) -> str:
        """Extracts text from an element, supporting hidden components and images (via markers)."""
        try:
            target_element = element
            if isinstance(element, tuple):
                target_element = self.find_element(element)
                
            if not isinstance(target_element, WebElement):
                raise TypeError("Input must be a (By, str) tuple or a WebElement.")

            if extract_hidden:
                raw_text = target_element.get_attribute("textContent")
                return raw_text.strip() if raw_text else ""
            else:
                # ========================================================
                # TEXT & IMAGE SCANNER
                # ========================================================
                js_script = """
                var target = arguments[0];
                
                function extractText(node) {
                    // STEP 1: PROCESS PURE TEXT NODES
                    if (node.nodeType === 3) { 
                        var parentStyle = window.getComputedStyle(node.parentNode);
                        if (parseFloat(parentStyle.fontSize) === 0 || 
                            parentStyle.display === 'none' || 
                            parentStyle.visibility === 'hidden' || 
                            parentStyle.opacity === '0' ||
                            parentStyle.color === 'rgba(0, 0, 0, 0)' ||
                            parentStyle.color === 'transparent') {
                            return "";
                        }
                        return node.nodeValue;
                    }
                    
                    // STEP 2: PROCESS ELEMENT NODES (div, p, span, img...)
                    if (node.nodeType === 1) { 
                        var style = window.getComputedStyle(node);
                        if (style.display === 'none' || style.visibility === 'hidden' || style.opacity === '0') {
                            return "";
                        }
                        
                        // ----------------------------------------------------
                        // IMAGE RECOGNITION AND PLACEHOLDER MARKING
                        // ----------------------------------------------------
                        if (node.tagName === 'IMG') {
                            var src = node.src || node.getAttribute('data-src') || node.getAttribute('data-original');
                            if (src) {
                                return "\\n[IMAGE_MARKER_START]" + src + "[IMAGE_MARKER_END]\\n";
                            }
                            return "";
                        }
                        // ----------------------------------------------------
                        
                        var text = "";
                        
                        // Handle line breaks <br> and block elements (p, div) for paragraph separation
                        if (node.tagName === 'BR' || style.display === 'block' || node.tagName === 'P') {
                            text += "\\n";
                        }
                        
                        // EXTRACT TEXT FROM CSS ::before
                        var before = window.getComputedStyle(node, '::before');
                        if (before && before.content && before.content !== 'none' && before.content !== 'normal') {
                            if (parseFloat(before.fontSize) > 0 && before.opacity !== '0') {
                                text += before.content.replace(/^["']|["']$/g, '');
                            }
                        }
                        
                        // Recurse into child nodes
                        for (var i = 0; i < node.childNodes.length; i++) {
                            text += extractText(node.childNodes[i]);
                        }
                        
                        // EXTRACT TEXT FROM CSS ::after
                        var after = window.getComputedStyle(node, '::after');
                        if (after && after.content && after.content !== 'none' && after.content !== 'normal') {
                            if (parseFloat(after.fontSize) > 0 && after.opacity !== '0') {
                                text += after.content.replace(/^["']|["']$/g, '');
                            }
                        }
                        
                        return text;
                    }
                    return "";
                }
                
                var result = extractText(target);
                
                // STEP 3: CLEAN TEXT
                result = result.replace(/\\r/g, '')
                               .replace(/[ \\t]+/g, ' ')
                               .replace(/\\n\\s+/g, '\\n')
                               .replace(/\\n{3,}/g, '\\n\\n');
                               
                return result.trim();
                """
                
                clean_text = self.driver.execute_script(js_script, target_element)
                return clean_text if clean_text else ""

        except StaleElementReferenceException:
            print("Warning: Element has been stale.")
            return ""
        except Exception as e:
            print(f"Unknown error: {e}")
            return ""

    def close_browser(self):
        """Closes the current tab."""
        self.driver.close()

    def quit_driver(self):
        """Quits the browser and ends the driver session."""
        self.driver.quit()

    def extract_images_to_file(self, elements: List[WebElement], file_name: str):
        """
        Scrolls to each image, scans with OCR, and saves to a Word file (.docx).
        """
        if not elements:
            print("Provided elements list is empty! Nothing to scan.")
            return

        print(f"Starting to scan {len(elements)} images and saving to: {file_name}")
        file_path = os.path.join(self.path, file_name + ".docx")

        # 1. Initialize Document (Create new or open existing)
        if os.path.exists(file_path):
            document = docx.Document(file_path)
        else:
            document = docx.Document()
            document.add_heading(f"OCR Data from Images: {file_name}", level=1)

        # 2. Loop through each image
        for index, img in enumerate(elements, start=1):
            try:
                # Scroll into view
                self.driver.execute_script("arguments[0].scrollIntoView({block: 'center'});", img)
                time.sleep(1.5) # Wait for Lazy-load
                
                # Capture screenshot and open with Pillow
                image_bytes = img.screenshot_as_png
                image = Image.open(io.BytesIO(image_bytes))
                
                # Call Tesseract
                print(f"Translating image {index}/{len(elements)}...")
                text = pytesseract.image_to_string(image, lang='vie')
                
                # Write to Document in RAM
                if text.strip():
                    document.add_paragraph(text.strip())
                    
            except Exception as e:
                print(f"Error at image {index}: {e}")
                document.add_paragraph(f"[Error reading image {index}]")

        # 3. Save all results from RAM to disk
        try:
            document.save(file_path)
            print(f"\nCOMPLETED! All content has been saved to {file_name}.docx")
        except Exception as e:
            print(f"Error saving DOCX file: {e}")
