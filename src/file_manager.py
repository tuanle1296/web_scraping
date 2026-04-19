import os
import json
import pathlib
import shutil
import re
import io
import requests
import docx
from docx.shared import Inches
from typing import Optional, Union, Any, List
from bs4 import BeautifulSoup
from selenium.webdriver.remote.webelement import WebElement


class FileManager:
    """Class to handle file and folder operations."""

    def __init__(self) -> None:
        self.path: Optional[str] = None

    def set_path(self, folder_name: str) -> None:
        """Specifies the working directory without creating it."""
        cur_dir = os.path.abspath(os.getcwd())
        self.path = os.path.join(cur_dir, folder_name)

    def create_folder(self, folder_name: str) -> None:
        """Creates a new folder and sets it as the current working path."""
        cur_dir = os.path.abspath(os.getcwd())
        dir_name = folder_name
        self.path = os.path.join(cur_dir, dir_name)
        pathlib.Path(self.path).mkdir(parents=True, exist_ok=True)

    @staticmethod
    def remove_file_if_exists(file: str) -> None:
        """Removes a file from the disk if it exists."""
        try:
            if os.path.exists(file):
                os.remove(file)
        except Exception:
            pass

    def zip_folder(self, folder_name: str) -> str:
        """Zips a folder and returns the path to the resulting zip file."""
        self.set_path(folder_name)
        if not self.path or not os.path.exists(self.path):
            raise FileNotFoundError(f"Folder not found: {self.path}")
            
        archive_path = shutil.make_archive(self.path, 'zip', self.path)
        print(f"Folder {self.path} zipped successfully to {archive_path}")
        return archive_path

    def define_img_name(self, title: str) -> str:
        """Defines the full path for an image file based on the title."""
        if not self.path:
            raise ValueError("Working path not set.")
        return os.path.join(self.path, title + ".png")

    def pass_data_to_file(self, source: str, file_name: str) -> None:
        """Writes data (string) to a file."""
        with open(file_name, "w") as f:
            f.write(source)

    def save_doc(self, title: Union[BeautifulSoup, WebElement, str], body: Union[BeautifulSoup, WebElement, str]) -> None:
        """Saves a title and body to a Word file (.docx)."""
        if not self.path:
            raise ValueError("Working path not set.")
        document = docx.Document()
        def get_text(obj: Any) -> str:
            if hasattr(obj, 'get_text'): return obj.get_text(separator='\n', strip=True)
            if hasattr(obj, 'text'): return obj.text
            return str(obj)
        title_text = get_text(title)
        document.add_heading(title_text)
        document.add_paragraph(get_text(body))
        document.save(os.path.join(self.path, f"{title_text}.docx"))

    def add_text_to_doc_file(self, title: str, text: str, file_name: Optional[str] = None) -> None:
        """Adds text (with image support via markers) to a new Word file."""
        if not self.path:
            raise ValueError("Working path not set.")
        document = docx.Document()
        try:
            document.add_heading(title)
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if not line:
                    document.add_paragraph(""); continue
                if "[IMAGE_MARKER_START]" in line:
                    parts = re.split(r'\[IMAGE_MARKER_START\](.*?)\[IMAGE_MARKER_END\]', line)
                    for part in parts:
                        part = part.strip()
                        if not part: continue
                        if part.startswith("http"):
                            try:
                                res = requests.get(part, timeout=10)
                                if res.status_code == 200:
                                    document.add_picture(io.BytesIO(res.content), width=Inches(5.0))
                            except Exception: pass
                        else: document.add_paragraph(part)
                else: document.add_paragraph(line)
            safe_name = file_name if file_name else re.sub(r'[\\/*?:"<>|]', "", title).strip()
            document.save(os.path.join(self.path, safe_name + ".docx"))
        except Exception as e:
            raise Exception(f"Word writing error: {e}")


def get_config_folder_id() -> Optional[str]:
    """Reads the Google Drive folder ID from config.json."""
    try:
        with open("config.json", "r") as f:
            return json.load(f).get("google_drive_folder_id")
    except Exception:
        return None
